import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import math
from PIL import Image, ImageDraw, ImageFont

# ---------------------------------------------------------
# CONSTANTS & DECORATION HELPERS
# ---------------------------------------------------------
COLOR_PRIMARY_HEX = "4C1D95"     # Deep Purple / Indigo
COLOR_SECONDARY_HEX = "7C3AED"   # Vibrant Purple / Violet
COLOR_DARK_HEX = "1F2937"        # Dark Gray (Body Text)
COLOR_LIGHT_BG_HEX = "F5F3FF"    # Very light purple background
COLOR_MUTED_BORDER_HEX = "D1D5DB"# Muted border color
COLOR_WARNING_HEX = "B91C1C"      # Crimson red for warnings/abuses

COLOR_PRIMARY = RGBColor(76, 29, 149)
COLOR_SECONDARY = RGBColor(124, 58, 237)
COLOR_DARK = RGBColor(31, 41, 55)
COLOR_MUTED = RGBColor(107, 114, 128)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    # dxa units: 20 dxa = 1 pt. 140 dxa = 7 pt, 180 dxa = 9 pt
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top="none", bottom="none", left="none", right="none", 
                     color="D1D5DB", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders_str = f'<w:tcBorders {nsdecls("w")}>'
    
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val == "none":
            borders_str += f'<w:{side} w:val="none"/>'
        else:
            borders_str += f'<w:{side} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            
    borders_str += '</w:tcBorders>'
    borders = parse_xml(borders_str)
    tcPr.append(borders)

def make_callout(doc, text, title="NOTE", level="info"):
    # Create a 1x1 table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    
    # Border & Shading based on level
    border_color = COLOR_PRIMARY_HEX
    bg_color = "F5F3FF"
    title_prefix = "💡 " + title
    
    if level == "warning":
        border_color = COLOR_WARNING_HEX
        bg_color = "FEF2F2"
        title_prefix = "⚠️ " + title
    elif level == "success":
        border_color = "059669" # Emerald
        bg_color = "ECFDF5"
        title_prefix = "✅ " + title
    
    set_cell_background(cell, bg_color)
    set_cell_left_border_thick(cell, border_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(title_prefix + "\n")
    r_title.bold = True
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(31, 41, 55)
    if level == "warning":
        r_title.font.color.rgb = RGBColor(185, 28, 28)
    elif level == "success":
        r_title.font.color.rgb = RGBColor(5, 150, 105)
        
    r_body = p.add_run(text)
    r_body.italic = True
    r_body.font.name = "Segoe UI"
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = RGBColor(75, 85, 99)
    
    # Add an empty spacing paragraph after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

def set_cell_left_border_thick(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{hex_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)

def set_table_light_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, top="single", bottom="single", left="none", right="none", color="E5E7EB", sz="4")

def add_heading_with_bottom_border(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    
    # Text run formatting
    run = h.runs[0]
    run.font.name = "Segoe UI Semibold"
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY
        # Add bottom border XML to paragraph
        pPr = h._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="{COLOR_PRIMARY_HEX}"/></w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_SECONDARY
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = COLOR_DARK
        
    return h

# ---------------------------------------------------------
# PROGRAMMATIC DIAGRAM GENERATORS (PIL)
# ---------------------------------------------------------
def get_segoe_fonts():
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/segoeuib.ttf", 16)
        font_body = ImageFont.truetype("C:/Windows/Fonts/Segoeui.ttf", 13)
        font_caption = ImageFont.truetype("C:/Windows/Fonts/segoeuib.ttf", 20)
        font_tag = ImageFont.truetype("C:/Windows/Fonts/Segoeui.ttf", 11)
    except:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()
        font_caption = ImageFont.load_default()
        font_tag = ImageFont.load_default()
    return font_title, font_body, font_caption, font_tag

def draw_arrow(draw, start, end, color, width=2):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    angle = math.atan2(dy, dx)
    wing_len = 10
    w1_x = x2 - wing_len * math.cos(angle - math.pi/6)
    w1_y = y2 - wing_len * math.sin(angle - math.pi/6)
    w2_x = x2 - wing_len * math.cos(angle + math.pi/6)
    w2_y = y2 - wing_len * math.sin(angle + math.pi/6)
    draw.polygon([(x2, y2), (w1_x, w1_y), (w2_x, w2_y)], fill=color)

def draw_card(draw, x, y, w, h, bg_color, border_color, title, text_lines, font_title, font_body):
    draw.rounded_rectangle([x, y, x+w, y+h], radius=8, fill=bg_color, outline=border_color, width=2)
    draw.text((x + 14, y + 10), title, fill=(255, 255, 255), font=font_title)
    curr_y = y + 36
    for line in text_lines:
        draw.text((x + 14, curr_y), line, fill=(199, 210, 254), font=font_body)
        curr_y += 18

def generate_diagram_architecture(output_path):
    img = Image.new("RGB", (1100, 520), "#1E1B4B") # Dark Indigo Canvas
    draw = ImageDraw.Draw(img)
    f_title, f_body, f_caption, f_tag = get_segoe_fonts()
    
    # Title
    draw.text((30, 20), "UNSEEN - System Architecture Architecture Overview", fill=(255, 255, 255), font=f_caption)
    draw.line([(30, 50), (1070, 50)], fill="#4F46E5", width=2)
    
    # LEVEL 1: CLIENTS
    draw_card(draw, 50, 90, 260, 110, "#312E81", "#6366F1", "Web Client (Next.js)", 
              ["- Tailwind CSS Theme", "- Socket.io Client Client", "- JWT session memory"], f_title, f_body)
              
    draw_card(draw, 50, 230, 260, 110, "#312E81", "#6366F1", "Mobile Client (Android APK)", 
              ["- Kotlin WebView wrapper", "- LocalStorage sessions", "- Built-in asset caches"], f_title, f_body)
              
    draw_card(draw, 50, 370, 260, 110, "#312E81", "#6366F1", "Admin Portal (Next.js)", 
              ["- Content Moderation UI", "- Report ticket review", "- User suspension dashboard"], f_title, f_body)

    # LEVEL 2: SERVERS (GATEWAY & APPLICATION)
    draw_card(draw, 420, 140, 280, 110, "#312E81", "#8B5CF6", "Express.js REST API Server", 
              ["- JWT Auth Middleware", "- Password hashing (Bcrypt)", "- Post & Interaction routers"], f_title, f_body)
              
    draw_card(draw, 420, 290, 280, 110, "#312E81", "#8B5CF6", "Socket.io WebSocket Server", 
              ["- Real-time chat pipes", "- Live online/typing status", "- Real-time badge sync updates"], f_title, f_body)

    # LEVEL 3: PERSISTENCE (DATABASES)
    draw_card(draw, 800, 210, 250, 150, "#312E81", "#EC4899", "MongoDB Atlas DB Layer", 
              ["- Users: Credentials & profiles", "- Posts & Comments feeds", "- Messages & Chat logs", "- Reports & suspension logs"], f_title, f_body)

    # CONNECTIONS
    # Client 1 -> REST API
    draw_arrow(draw, (310, 145), (420, 185), "#10B981", 2)
    # Client 1 -> Sockets
    draw_arrow(draw, (310, 155), (420, 320), "#10B981", 2)
    
    # Client 2 -> REST API
    draw_arrow(draw, (310, 275), (420, 205), "#10B981", 2)
    # Client 2 -> Sockets
    draw_arrow(draw, (310, 285), (420, 345), "#10B981", 2)
    
    # Client 3 -> REST API
    draw_arrow(draw, (310, 425), (420, 225), "#10B981", 2)

    # REST API -> MongoDB
    draw_arrow(draw, (700, 195), (800, 250), "#06B6D4", 2)
    # Sockets -> MongoDB
    draw_arrow(draw, (700, 345), (800, 310), "#06B6D4", 2)

    # Labels on arrows
    draw.text((325, 120), "HTTP/HTTPS", fill="#A5B4FC", font=f_tag)
    draw.text((325, 335), "WebSockets", fill="#A5B4FC", font=f_tag)
    draw.text((715, 205), "Mongoose", fill="#A5B4FC", font=f_tag)
    
    img.save(output_path, "PNG")

def generate_diagram_registration(output_path):
    img = Image.new("RGB", (1100, 240), "#1E1B4B")
    draw = ImageDraw.Draw(img)
    f_title, f_body, f_caption, f_tag = get_segoe_fonts()
    
    # Title
    draw.text((30, 20), "Registration & Identity Generation Flow", fill=(255, 255, 255), font=f_caption)
    draw.line([(30, 50), (1070, 50)], fill="#4F46E5", width=2)
    
    # Step Cards
    draw_card(draw, 30, 80, 180, 110, "#312E81", "#6366F1", "1. User Signup", 
              ["- Inputs display nickname", "- Inputs credentials", "- Optional Recovery Email"], f_title, f_body)
              
    draw_card(draw, 240, 80, 200, 110, "#312E81", "#8B5CF6", "2. Anonymity Engine", 
              ["- Generates random name", "- Combines Adj + Noun", "- Example: 'silent-wolf-22'"], f_title, f_body)
              
    draw_card(draw, 470, 80, 190, 110, "#312E81", "#8B5CF6", "3. Cryptography", 
              ["- Salt + Hash password", "- Standard 10-round bcrypt", "- Pre-save hooks trigger"], f_title, f_body)

    draw_card(draw, 690, 80, 180, 110, "#312E81", "#EC4899", "4. MongoDB Atlas", 
              ["- Saves IUser record", "- Initial trustScore = 70", "- Index checks pass"], f_title, f_body)

    draw_card(draw, 900, 80, 170, 110, "#312E81", "#10B981", "5. Session Load", 
              ["- Server issues JWT", "- Frontend saves token", "- Auto-connects socket"], f_title, f_body)

    # Connections
    draw_arrow(draw, (210, 135), (240, 135), "#06B6D4", 2)
    draw_arrow(draw, (440, 135), (470, 135), "#06B6D4", 2)
    draw_arrow(draw, (660, 135), (690, 135), "#06B6D4", 2)
    draw_arrow(draw, (870, 135), (900, 135), "#06B6D4", 2)

    img.save(output_path, "PNG")

def generate_diagram_messaging(output_path):
    img = Image.new("RGB", (1100, 480), "#1E1B4B")
    draw = ImageDraw.Draw(img)
    f_title, f_body, f_caption, f_tag = get_segoe_fonts()
    
    # Title
    draw.text((30, 20), "Real-Time Messaging & Dynamic Read Badge Synchronization", fill=(255, 255, 255), font=f_caption)
    draw.line([(30, 50), (1070, 50)], fill="#4F46E5", width=2)
    
    # Nodes
    draw_card(draw, 50, 90, 220, 120, "#312E81", "#6366F1", "Sender (Client A)", 
              ["- Sends message", "- Body: text, receiverId", "- Fires socket event", "  'send_message'"], f_title, f_body)
              
    draw_card(draw, 420, 90, 260, 130, "#312E81", "#8B5CF6", "Socket.io & Node Server", 
              ["- Verifies sender token", "- Writes message to DB", "- If Receiver ONLINE:", "  emits 'new_message'", "- Else: increments badge"], f_title, f_body)
              
    draw_card(draw, 830, 90, 220, 120, "#312E81", "#6366F1", "Receiver (Client B)", 
              ["- Receives 'new_message'", "- Displays chat bubble", "- Plays notify sound", "- Sends back 'mark_read'"], f_title, f_body)

    # Sync Node for Badges
    draw_card(draw, 420, 310, 260, 120, "#312E81", "#EC4899", "Offline Badge Synchronizer", 
              ["- User B reconnects to network", "- Queries unread count in DB", "- Server emits 'unread_badge_sync'", "- Client B updates nav badges"], f_title, f_body)

    # Connections
    # A -> Server
    draw_arrow(draw, (270, 140), (420, 140), "#10B981", 2)
    # Server -> B
    draw_arrow(draw, (680, 140), (830, 140), "#10B981", 2)
    # B -> Server (mark_read)
    draw_arrow(draw, (830, 180), (680, 180), "#F59E0B", 2)
    
    # DB Sync path (Server to Sync)
    draw_arrow(draw, (550, 220), (550, 310), "#06B6D4", 2)
    # Sync -> Receiver B
    draw_arrow(draw, (680, 370), (830, 190), "#EC4899", 2)

    # Text descriptions on paths
    draw.text((290, 115), "WS 'send_message'", fill="#A5B4FC", font=f_tag)
    draw.text((700, 115), "WS 'new_message'", fill="#A5B4FC", font=f_tag)
    draw.text((700, 185), "WS 'mark_read'", fill="#FBCFE8", font=f_tag)
    draw.text((710, 340), "WS 'unread_badge_sync'", fill="#FBCFE8", font=f_tag)

    img.save(output_path, "PNG")

def generate_diagram_moderation(output_path):
    img = Image.new("RGB", (1100, 380), "#1E1B4B")
    draw = ImageDraw.Draw(img)
    f_title, f_body, f_caption, f_tag = get_segoe_fonts()
    
    # Title
    draw.text((30, 20), "Abuse Reporting & Moderation State Machine", fill=(255, 255, 255), font=f_caption)
    draw.line([(30, 50), (1070, 50)], fill="#4F46E5", width=2)
    
    # States
    draw_card(draw, 50, 120, 200, 110, "#312E81", "#10B981", "Active Post State", 
              ["- Created by anonymous user", "- Visible on global feed", "- trustScore >= 70"], f_title, f_body)
              
    draw_card(draw, 350, 120, 220, 110, "#312E81", "#F59E0B", "Flagged & Logged State", 
              ["- User clicks 'Report'", "- API inserts 'Report' entry", "- reportsCount increments"], f_title, f_body)
              
    draw_card(draw, 670, 120, 220, 110, "#312E81", "#EF4444", "Auto-Escalated State", 
              ["- Reports count reaches 5+", "- Post auto-hidden from feed", "- Poster trustScore drops by 10"], f_title, f_body)

    # Resolution Nodes (split branch)
    draw_card(draw, 930, 70, 150, 90, "#1F2937", "#10B981", "A. Dismissed", 
              ["- Report set resolved", "- Post restored", "- Reporter penalized"], f_title, f_body)
              
    draw_card(draw, 930, 190, 150, 110, "#451A03", "#EF4444", "B. Suspended", 
              ["- Post permanently deleted", "- Suspension logged", "- Author lockout check"], f_title, f_body)

    # Connections
    # Active -> Flagged
    draw_arrow(draw, (250, 175), (350, 175), "#06B6D4", 2)
    # Flagged -> Escalated
    draw_arrow(draw, (570, 175), (670, 175), "#06B6D4", 2)
    # Escalated -> Dismissed
    draw_arrow(draw, (890, 160), (930, 110), "#10B981", 2)
    # Escalated -> Suspended
    draw_arrow(draw, (890, 190), (930, 230), "#EF4444", 2)

    # Condition labels
    draw.text((260, 150), "Report Action", fill="#A5B4FC", font=f_tag)
    draw.text((585, 150), "Reports >= 5", fill="#EF4444", font=f_tag)
    draw.text((850, 115), "Legitimate", fill="#10B981", font=f_tag)
    draw.text((850, 220), "Violative", fill="#EF4444", font=f_tag)

    img.save(output_path, "PNG")

# ---------------------------------------------------------
# WORD DOCUMENT GENERATOR (.DOCX)
# ---------------------------------------------------------
def build_word_document():
    doc = docx.Document()
    
    # Page setup - Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default style to Segoe UI
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = COLOR_DARK
    
    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    # Empty paragraph for top spacing
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(80)
    
    # App Logo Title
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("🌌 UNSEEN")
    run_logo.font.name = "Segoe UI Semibold"
    run_logo.font.size = Pt(32)
    run_logo.font.color.rgb = COLOR_PRIMARY
    run_logo.bold = True
    
    p_sublogo = doc.add_paragraph()
    p_sublogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sublogo.paragraph_format.space_after = Pt(40)
    run_sublogo = p_sublogo.add_run("Anonymous. Secure. Unfiltered.")
    run_sublogo.font.name = "Segoe UI"
    run_sublogo.font.size = Pt(12)
    run_sublogo.font.color.rgb = COLOR_SECONDARY
    run_sublogo.italic = True
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("SYSTEM ARCHITECTURE & PRODUCT DOCUMENTATION")
    run_title.font.name = "Segoe UI Semibold"
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = COLOR_DARK
    run_title.bold = True
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_after = Pt(140)
    run_subtitle = p_subtitle.add_run("A Unified Guide containing Product Requirements (PRD), Technical Requirements (TRD), and Core Flow Workflows.")
    run_subtitle.font.name = "Segoe UI"
    run_subtitle.font.size = Pt(11)
    run_subtitle.font.color.rgb = COLOR_DARK
    
    # Author Block
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(6)
    run_auth_lbl = p_author.add_run("Prepared by:\n")
    run_auth_lbl.font.size = Pt(9.5)
    run_auth_lbl.font.color.rgb = COLOR_MUTED
    
    run_author = p_author.add_run("Krishnam Gupta\n")
    run_author.font.name = "Segoe UI Semibold"
    run_author.font.size = Pt(12)
    run_author.font.color.rgb = COLOR_PRIMARY
    run_author.bold = True
    
    run_details = p_author.add_run("B.Tech CSE (AI/ML) Student | GLA University, Mathura")
    run_details.font.size = Pt(10)
    run_details.font.color.rgb = COLOR_DARK
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(6)
    run_date = p_date.add_run("Date: June 21, 2026")
    run_date.font.size = Pt(9.5)
    run_date.font.color.rgb = COLOR_MUTED
    
    # Page Break
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # SECTION 1: PRODUCT REQUIREMENTS DOCUMENT (PRD)
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Section 1: Product Requirements Document (PRD)", level=1)
    
    p_intro = doc.add_paragraph(
        "UNSEEN is a digital ecosystem tailored for individuals seeking authentic self-expression. In modern social spaces, "
        "reputation pressure, professional branding, and social circles restrict people from expressing their rawest emotions, "
        "confessions, and vulnerabilities. UNSEEN dismantles these barriers by enforcing anonymity as a core protocol, "
        "while providing a real-time responsive social loop equivalent to mainstream networks."
    )
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.paragraph_format.space_after = Pt(10)
    
    add_heading_with_bottom_border(doc, "1.1 Vision & Key Objectives", level=2)
    
    # Bullets
    bullets = [
        ("🔐 Privacy-First Identity System:", " Identity is protected using a dual-name structure (a private login email/password and an auto-generated random username combination like 'silent-shadow-82'). Display names are customizable, but usernames cannot link to real-world identifiers."),
        ("⚡ High-Fidelity Real-Time Interactions:", " Post creation, comment feeds, likes, saves, and direct messaging are processed instantaneously, utilizing low-latency server connections and dynamic client updates."),
        ("🛡️ Safe & Clean Community Ecosystem:", " Anonymity does not excuse abuse. A robust reporting system auto-flags violative content, dynamically adjusting user trust scores to isolate trolls while safeguarding genuine confiders."),
        ("📱 Universal Device Access:", " Mobile-first responsiveness paired with a native-wrapped Android APK client ensures seamless interaction across desktop web and mobile screens.")
    ]
    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.line_spacing = 1.15
        
        run_t = bp.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = COLOR_PRIMARY
        bp.add_run(desc)
        
    add_heading_with_bottom_border(doc, "1.2 User Personas", level=2)
    
    personas = [
        ("Aarav (The Confessor - Age 21)", "Needs a therapeutic outlet to vent academic pressure, doubts, and anxiety without risking social repercussions. Valuing extreme privacy, Aarav utilizes UNSEEN's auto-generated username engine to confess details he would never discuss publicly."),
        ("Elena (The Supporter - Age 25)", "Enjoys reading genuine human stories and providing empathetic counsel. She frequently uses the private chat channels to offer support, using message reactions and fast direct text to keep the conversations active."),
        ("Sam (The Administrator/Moderator - Age 29)", "Responsible for keeping the environment safe. Sam relies on automated report aggregation dashboards to review content flagged by users, suspend bad actors, and preserve the system's safe-space parameters.")
    ]
    for title, desc in personas:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.line_spacing = 1.15
        
        run_t = bp.add_run(title + ": ")
        run_t.bold = True
        run_t.font.color.rgb = COLOR_SECONDARY
        bp.add_run(desc)

    add_heading_with_bottom_border(doc, "1.3 Functional Feature Matrix", level=2)
    
    # Table of Features
    table = doc.add_table(rows=8, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(0.8), Inches(1.5), Inches(2.2), Inches(0.8), Inches(1.2)]
    headers = ["ID", "Feature Module", "Functional Requirement Description", "Priority", "Status"]
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
    
    features_data = [
        ("F-01", "Identity Engine", "Automatic generation of unique, non-identifiable combinations of adjectives and nouns on user signup. Avatar color sets default custom styling.", "P0", "Implemented"),
        ("F-02", "Anonymous Feed", "Post creation support with real-time global feed loading, post interactions (likes, saves, delete capabilities), and comments threads.", "P0", "Implemented"),
        ("F-03", "Real-Time Chat", "One-to-one direct messages powered by Socket.io, support for emoji reactions, unread message badges, and typing indicators.", "P0", "Implemented"),
        ("F-04", "Social Network", "Follow/unfollow system for individual creators, user bio customization, search bar query matches, and personal feeds.", "P1", "Implemented"),
        ("F-05", "Report Moderation", "Report button on posts/comments. Aggregates reports, automatically hides posts reaching 5+ flags, and updates poster's trustScore.", "P0", "Implemented"),
        ("F-06", "Security Lockouts", "Throttles failed OTP password reset attempts and logins, locking accounts after 5 failures to protect against brute-force attacks.", "P1", "Implemented"),
        ("F-07", "Native Android", "APK distribution wrapping the responsive web application inside a lightweight native Android wrapper with adaptive launcher icons.", "P0", "Implemented")
    ]
    
    for row_idx, row_data in enumerate(features_data, start=1):
        row_cells = table.rows[row_idx].cells
        # Alternating row bg
        row_bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            # Styling cell
            set_cell_background(row_cells[col_idx], row_bg)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.0)
                p.runs[0].font.color.rgb = COLOR_DARK
                if col_idx == 0:
                    p.runs[0].bold = True
                if col_idx == 3: # Priority
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if text == "P0":
                        p.runs[0].bold = True
                        p.runs[0].font.color.rgb = COLOR_PRIMARY
                if col_idx == 4: # Status
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if text == "Implemented":
                        p.runs[0].bold = True
                        p.runs[0].font.color.rgb = RGBColor(5, 150, 105)
                        
    # Set widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    set_table_light_borders(table)
    doc.add_paragraph() # spacing
    
    add_heading_with_bottom_border(doc, "1.4 Non-Functional Requirements (NFRs)", level=2)
    
    nfrs = [
        ("🔐 Cryptographic Security:", " Passwords must be hashed using Bcrypt before DB write. Session validation relies on secure JWT tokens. Active sessions must lock out after 5 consecutive failures to prevent automated logins."),
        ("🛡️ Privacy Compliance:", " Zero personally identifiable information (PII) should leak in public endpoints. Email is encrypted at rest and is used strictly for authentication recovery, never linked on public profiles."),
        ("⚡ Low Latency Updates:", " Real-time websocket actions (chat delivery, read sync, typing indications) must resolve across active connections in under 120ms to avoid UI stutter."),
        ("🌐 Cross-Platform Compatibility:", " Complete visual responsiveness on desktop, tablet, and mobile browsers, alongside a highly optimized Kotlin-wrapped Android APK (retaining UI scaling and icon masks).")
    ]
    for title, desc in nfrs:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.line_spacing = 1.15
        
        run_t = bp.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = COLOR_PRIMARY
        bp.add_run(desc)

    add_heading_with_bottom_border(doc, "1.5 Future Roadmap & Enhancements", level=2)
    
    roadmap = [
        ("🎙️ Voice Confessions:", " Allowing users to post temporary audio logs with optional voice-changer filters to further secure verbal anonymity."),
        ("👥 Community Circles:", " Sub-communities created around specific interests (e.g., career vent, relationships, mental wellness) governed by community rules."),
        ("🤖 AI Moderation Engine:", " Deploying NLP toxicity classification models on the Express server to flag hate speech and self-harm keywords before they populate the global feed.")
    ]
    for title, desc in roadmap:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.line_spacing = 1.15
        
        run_t = bp.add_run(title)
        run_t.bold = True
        run_t.font.color.rgb = COLOR_SECONDARY
        bp.add_run(desc)

    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 2: TECHNICAL REQUIREMENTS DOCUMENT (TRD)
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Section 2: Technical Requirements Document (TRD)", level=1)
    
    doc.add_paragraph(
        "This section details the tech stack, structural models, API routing definitions, "
        "and security frameworks governing the UNSEEN deployment. The platform operates on a "
        "decoupled Client-Server architecture utilizing a real-time event-driven connection pipeline."
    )
    
    add_heading_with_bottom_border(doc, "2.1 Tech Stack & Infrastructure", level=2)
    
    # Table of Tech Stack
    table_stack = doc.add_table(rows=7, cols=4)
    table_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_stack.autofit = False
    
    col_widths_stack = [Inches(1.5), Inches(1.8), Inches(1.5), Inches(1.7)]
    headers_stack = ["Architecture Layer", "Core Frameworks / Tools", "Deployment Target", "Primary Purpose"]
    
    hdr_cells = table_stack.rows[0].cells
    for i, header in enumerate(headers_stack):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], COLOR_SECONDARY_HEX)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        
    stack_data = [
        ("Web Frontend", "React, Next.js (App Router), Tailwind CSS, Framer Motion", "Vercel Hosting", "Delivers responsive, animated web interface with fast Server-Side Render pipelines."),
        ("Android Wrapper", "Kotlin, Android Native SDK, WebView, XML Assets", "Google Play / APK Distribution", "Wraps responsive web portal in a native client container with custom icon scaling."),
        ("Backend Services", "Node.js, Express.js API, TypeScript, Cors, Helmet", "Render (Web Service)", "Handles REST API routing, user registration validation, security, and OTP operations."),
        ("Real-time Server", "Socket.io Engine (WebSockets), Node.js cluster", "Render (Shared instance)", "Orchestrates low-latency messaging, active connection states, and read indicator sync."),
        ("Database Layer", "MongoDB Atlas Cloud, Mongoose Object Relational Mapper", "MongoDB Atlas Managed", "Hosts document collection clusters, checks indexes, and runs analytical aggregations."),
        ("Cryptographic Engine", "Bcrypt Hashing, JWT Auth Tokens, Crypto OTP API", "Server Container runtime", "Locks passwords via 10-round salt hashes and signs cryptographically secure access keys.")
    ]
    
    for row_idx, row_data in enumerate(stack_data, start=1):
        row_cells = table_stack.rows[row_idx].cells
        row_bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], row_bg)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.0)
                p.runs[0].font.color.rgb = COLOR_DARK
                if col_idx == 0:
                    p.runs[0].bold = True
                    
    for row in table_stack.rows:
        for idx, width in enumerate(col_widths_stack):
            row.cells[idx].width = width
            
    set_table_light_borders(table_stack)
    doc.add_paragraph()
    
    add_heading_with_bottom_border(doc, "2.2 Database Schemas & Data Model Reference", level=2)
    
    doc.add_paragraph("UNSEEN stores records inside MongoDB. Relationships between entities are managed using Mongoose ObjectIds.")
    
    # Model 1: User
    p_model = doc.add_paragraph()
    p_model.paragraph_format.space_before = Pt(8)
    p_model.paragraph_format.space_after = Pt(2)
    run_model = p_model.add_run("👤 Model: User (Collection: users)")
    run_model.bold = True
    run_model.font.size = Pt(11)
    run_model.font.color.rgb = COLOR_PRIMARY
    
    table_m1 = doc.add_table(rows=6, cols=3)
    table_m1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_m1.autofit = False
    
    hdr_cells = table_m1.rows[0].cells
    hdr_cells[0].text = "Schema Key"
    hdr_cells[1].text = "Type / Config"
    hdr_cells[2].text = "Description & Purpose"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.0)
        set_cell_background(cell, "4B5563") # Slate-600
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        
    m1_data = [
        ("username", "String (Unique, Indexed, Lowercase)", "Auto-generated random persona name (e.g., 'mystic-breeze-47'). Public key."),
        ("displayName", "String (Max 50 chars, Required)", "Customizable display name created by the user for feed customization."),
        ("email", "String (Unique, Sparse Index, Optional)", "Used only for secure OTP account recovery. Protected from public exposure."),
        ("passwordHash", "String (Required)", "Stored securely using Bcrypt (10 rounds). Verified on REST logins."),
        ("trustScore", "Number (Default: 70, Min: 0, Max: 100)", "Quantifies community behavior. Drops on reporting, blocks account if low.")
    ]
    for row_idx, row_data in enumerate(m1_data, start=1):
        row_cells = table_m1.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.0)
            if col_idx == 0:
                row_cells[col_idx].paragraphs[0].runs[0].bold = True
                
    set_table_light_borders(table_m1)
    doc.add_paragraph()

    # Model 2: Post
    p_model = doc.add_paragraph()
    p_model.paragraph_format.space_before = Pt(8)
    p_model.paragraph_format.space_after = Pt(2)
    run_model = p_model.add_run("📝 Model: Post (Collection: posts)")
    run_model.bold = True
    run_model.font.size = Pt(11)
    run_model.font.color.rgb = COLOR_PRIMARY
    
    table_m2 = doc.add_table(rows=5, cols=3)
    table_m2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_m2.autofit = False
    
    hdr_cells = table_m2.rows[0].cells
    hdr_cells[0].text = "Schema Key"
    hdr_cells[1].text = "Type / Config"
    hdr_cells[2].text = "Description & Purpose"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.0)
        set_cell_background(cell, "4B5563")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        
    m2_data = [
        ("author", "Schema.Types.ObjectId (Ref: User)", "Reference to the User document that created the post. Anonymous in endpoints."),
        ("content", "String (Max 500 chars, Required)", "Text content representing the confession, feeling, or thought."),
        ("likes / saves", "Array [ ObjectId ]", "Tracks active user interactions. Used to construct personal saves grids."),
        ("reportsCount", "Number (Default: 0)", "Accumulates user abuse flags. Post is hidden from feed if counts exceed 5.")
    ]
    for row_idx, row_data in enumerate(m2_data, start=1):
        row_cells = table_m2.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.0)
            if col_idx == 0:
                row_cells[col_idx].paragraphs[0].runs[0].bold = True
                
    set_table_light_borders(table_m2)
    doc.add_paragraph()

    # Model 3: Message
    p_model = doc.add_paragraph()
    p_model.paragraph_format.space_before = Pt(8)
    p_model.paragraph_format.space_after = Pt(2)
    run_model = p_model.add_run("💬 Model: Message (Collection: messages)")
    run_model.bold = True
    run_model.font.size = Pt(11)
    run_model.font.color.rgb = COLOR_PRIMARY
    
    table_m3 = doc.add_table(rows=6, cols=3)
    table_m3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_m3.autofit = False
    
    hdr_cells = table_m3.rows[0].cells
    hdr_cells[0].text = "Schema Key"
    hdr_cells[1].text = "Type / Config"
    hdr_cells[2].text = "Description & Purpose"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.0)
        set_cell_background(cell, "4B5563")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        
    m3_data = [
        ("sender / receiver", "Schema.Types.ObjectId (Ref: User)", "Mapping representing the sender and recipient of the direct message."),
        ("text", "String (Required, Trimmed)", "Message content. Encrypted in transport, stored in clear DB documents."),
        ("reactions", "Array [ { user: ObjectId, emoji: String } ]", "Allows receivers/senders to stamp emoji reactions to the message."),
        ("isRead", "Boolean (Default: false)", "Tracks unread message states. Used to calculate real-time navigation notification dots."),
        ("createdAt", "Date", "System timestamp used to order conversation history chronologically.")
    ]
    for row_idx, row_data in enumerate(m3_data, start=1):
        row_cells = table_m3.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.0)
            if col_idx == 0:
                row_cells[col_idx].paragraphs[0].runs[0].bold = True
                
    set_table_light_borders(table_m3)
    doc.add_paragraph()

    add_heading_with_bottom_border(doc, "2.3 API REST Definitions", level=2)
    
    # Table of API endpoints
    table_api = doc.add_table(rows=7, cols=4)
    table_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_api.autofit = False
    
    col_widths_api = [Inches(0.8), Inches(1.8), Inches(1.8), Inches(2.1)]
    headers_api = ["Method", "Route Path", "Payload Parameters", "Expected JSON Response"]
    
    hdr_cells = table_api.rows[0].cells
    for i, header in enumerate(headers_api):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        
    api_data = [
        ("POST", "/api/auth/signup", "displayName, password, email (opt)", "{ token, user: { username, displayName } }"),
        ("POST", "/api/auth/login", "username, password", "{ token, user: { username, displayName, avatarColor } }"),
        ("POST", "/api/posts", "content, tags (opt) [Header: JWT]", "{ success: true, post: { id, content, author, reportsCount } }"),
        ("GET", "/api/posts", "[Header: JWT] query: page, limit", "{ posts: [ { id, content, likesCount, saved } ], hasMore: boolean }"),
        ("POST", "/api/posts/:id/report", "reason, details (opt) [Header: JWT]", "{ success: true, message: 'Report logged successfully' }"),
        ("GET", "/api/messages/unread-count", "[Header: JWT]", "{ unreadMessagesCount: number }")
    ]
    
    for row_idx, row_data in enumerate(api_data, start=1):
        row_cells = table_api.rows[row_idx].cells
        row_bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], row_bg)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.0)
                p.runs[0].font.color.rgb = COLOR_DARK
                if col_idx == 0:
                    p.runs[0].bold = True
                    if text == "POST":
                        p.runs[0].font.color.rgb = RGBColor(5, 150, 105)
                    elif text == "GET":
                        p.runs[0].font.color.rgb = COLOR_SECONDARY
                        
    for row in table_api.rows:
        for idx, width in enumerate(col_widths_api):
            row.cells[idx].width = width
            
    set_table_light_borders(table_api)
    doc.add_paragraph()

    add_heading_with_bottom_border(doc, "2.4 WebSockets Protocol Specifications", level=2)
    
    doc.add_paragraph("Real-time messaging channels use Socket.io. Clients connect using their signed JWT payload.")
    
    # WebSocket table
    table_ws = doc.add_table(rows=6, cols=3)
    table_ws.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ws.autofit = False
    
    col_widths_ws = [Inches(1.8), Inches(1.8), Inches(2.9)]
    headers_ws = ["Socket Event Name", "Direction / Source", "Event Payload & UI Sync Hook"]
    
    hdr_cells = table_ws.rows[0].cells
    for i, header in enumerate(headers_ws):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], COLOR_SECONDARY_HEX)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        
    ws_data = [
        ("join_room", "Client -> Server", "Payload: { conversationId, userId }\nAction: Registers client socket context into a dedicated chat room."),
        ("send_message", "Client -> Server", "Payload: { senderId, receiverId, text, conversationId }\nAction: Persists document in DB and broadcasts message to room."),
        ("new_message", "Server -> Client", "Payload: { id, senderId, receiverId, text, createdAt, isRead }\nAction: Client app adds bubble to UI and plays notification chime."),
        ("mark_read", "Client -> Server", "Payload: { messageId, receiverId, conversationId }\nAction: Updates message in DB to isRead=true. Syncs sender screen."),
        ("unread_badge_sync", "Server -> Client", "Payload: { totalUnreadCount: number }\nAction: Directly updates navigation message indicators (critical fix).")
    ]
    
    for row_idx, row_data in enumerate(ws_data, start=1):
        row_cells = table_ws.rows[row_idx].cells
        row_bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], row_bg)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.0)
                p.runs[0].font.color.rgb = COLOR_DARK
                if col_idx == 0:
                    p.runs[0].bold = True
                    
    for row in table_ws.rows:
        for idx, width in enumerate(col_widths_ws):
            row.cells[idx].width = width
            
    set_table_light_borders(table_ws)
    
    # Spacer
    doc.add_paragraph()
    
    # Callout about security
    make_callout(doc, 
                 "WebSocket connections undergo authorization checks during socket handshake. "
                 "If the JWT token expires or is malformed, the connection is instantly closed. "
                 "The server enforces strict connection indexing to tie socket IDs to active database User sessions.",
                 title="WEBSOCKET SECURITY RULE", level="success")

    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 3: SYSTEM WORKFLOWS & LIFE CYCLES
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "Section 3: System Workflows & Lifecycle Diagrams", level=1)
    
    doc.add_paragraph(
        "This section illustrates the logical flows running behind UNSEEN, "
        "visualizing how client requests and data transitions behave inside the system."
    )
    
    # Workflow 1: Architecture
    add_heading_with_bottom_border(doc, "3.1 System Architecture Flow", level=2)
    doc.add_paragraph(
        "The system separates user interfaces from backend operations. Next.js and Kotlin "
        "clients use secure API routing for posting, searching, and registration, while relying "
        "on WebSockets for active chat communication. MongoDB handles indexing and raw storage."
    )
    doc.add_picture("diagram_architecture.png", width=Inches(6.0))
    doc.add_paragraph("Figure 1: High-level visual system topology, showcasing decoupled layer connectivity.")
    doc.add_paragraph()

    # Workflow 2: Registration
    add_heading_with_bottom_border(doc, "3.2 Anonymous Registration & Identity Generation Flow", level=2)
    doc.add_paragraph(
        "UNSEEN intercepts regular signup routines to enforce privacy. When a user creates an account, "
        "an Anonymity Engine generates a random username by pairing descriptive adjectives with nouns, "
        "appending random digits to ensure absolute uniqueness. The password is encrypted before database storage."
    )
    doc.add_picture("diagram_registration.png", width=Inches(6.0))
    doc.add_paragraph("Figure 2: Step-by-step registration flow showing automatic random username generation.")
    doc.add_paragraph()

    # Workflow 3: Chat Sync
    add_heading_with_bottom_border(doc, "3.3 Real-Time Messaging & Read Badge Sync Lifecycle", level=2)
    doc.add_paragraph(
        "Real-time chatting uses a dedicated room system. When a sender transmits a message, "
        "the WebSocket server saves it and verifies the recipient's state. If online, the server "
        "emits a new message event immediately. If offline, the unread count increments, which "
        "is synchronized to the receiver's layout the moment they reconnect."
    )
    doc.add_picture("diagram_messaging.png", width=Inches(6.0))
    doc.add_paragraph("Figure 3: Sequence representing active messaging pipelines and dynamic read badge updates.")
    doc.add_paragraph()

    # Workflow 4: Moderation
    add_heading_with_bottom_border(doc, "3.4 Abuse Reporting & Content Moderation Lifecycle", level=2)
    doc.add_paragraph(
        "To prevent abuse, UNSEEN features a community report system. If a post receives "
        "multiple reports, the backend automatically flags the post, decrements the poster's "
        "trust score, and removes it from public feeds. Moderators review the queue to either "
        "delete the post permanently or dismiss the reports."
    )
    doc.add_picture("diagram_moderation.png", width=Inches(6.0))
    doc.add_paragraph("Figure 4: State transitions of reported content under the moderation guidelines.")
    doc.add_paragraph()

    # Callout Warning
    make_callout(doc, 
                 "A user whose trustScore falls below 40 points is marked as suspended (isSuspended: true). "
                 "Any active login token is revoked, and the user's IP is temporarily throttled on the gateway "
                 "to prevent automated signup re-runs.", 
                 title="MODERATION CONSTRAINTS WARNING", level="warning")

    # Save
    doc.save("UNSEEN_Product_Technical_Documentation.docx")
    print("UNSEEN_Product_Technical_Documentation.docx generated successfully!")

# ---------------------------------------------------------
# SCRIPT ENTRYPOINT
# ---------------------------------------------------------
if __name__ == "__main__":
    print("Step 1: Generating diagram PNG assets...")
    generate_diagram_architecture("diagram_architecture.png")
    generate_diagram_registration("diagram_registration.png")
    generate_diagram_messaging("diagram_messaging.png")
    generate_diagram_moderation("diagram_moderation.png")
    print("Diagrams created.")

    print("Step 2: Compiling Word document...")
    build_word_document()

    print("Step 3: Cleaning up temporary PNG assets...")
    try:
        os.remove("diagram_architecture.png")
        os.remove("diagram_registration.png")
        os.remove("diagram_messaging.png")
        os.remove("diagram_moderation.png")
        print("Temporary files removed successfully.")
    except Exception as e:
        print(f"Error during cleanup: {e}")

    print("Generation process complete!")
